import pandas as pd
import mysql.connector
from openpyxl import Workbook
from openpyxl.styles import Alignment
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Database connection parameters
db_config = {
    'host': '127.0.0.1',
    'user': 'root',
    'password': '',
    'database': 'schedopt_db'
}

# Mapping of program acronyms to full names
PROGRAM_NAMES = {
    'AEET': 'Associate in Electrical Engineering Technology',
    'BACOMM': 'Bachelor of Arts in Communication',
    'BAELS': 'Bachelor of Arts in English Language Studies',
    'BAINDIS': 'Bachelor of Arts in Interdisciplinary Studies',
    'BAINTS': 'Bachelor of Arts in International Studies',
    'BAPHILO': 'Bachelor of Arts in Philosophy',
    'BEED': 'Bachelor of Elementary Education',
    'BPED': 'Bachelor of Physical Education',
    'BSBME': 'Bachelor of Science in Biomedical Engineering',
    'BSCE': 'Bachelor of Science in Civil Engineering',
    'BSCPE': 'Bachelor of Science in Computer Engineering',
    'BSCS': 'Bachelor of Science in Computer Science',
    'BSECE': 'Bachelor of Science in Electronics and Communications Engineering',
    'BSIT': 'Bachelor of Science in Information Technology',
    'BSLM': 'Bachelor of Science in Legal Management',
    'BSMATH': 'Bachelor of Science in Mathematics',
    'BSN': 'Bachelor of Science in Nursing',
    'BSNMCA': 'Bachelor of Science in New Media and Computer Animation',
    'BSOA': 'Bachelor of Science in Office Administration',
    'CON': 'College of Nursing',
    'CSITE': 'College of Science, Information Technology and Engineering',
    'SED': 'School of Education',
    'SLA': 'School of Liberal Arts',
    'SMA': 'School of Management and Accountancy'
}

def get_data_from_db():
    """Fetch data from the tbl_final_assignment table"""
    try:
        conn = mysql.connector.connect(**db_config)
        query = "SELECT * FROM tbl_final_assignment"
        df = pd.read_sql(query, conn)
        conn.close()
        return df
    except Exception as e:
        print(f"Error fetching data from database: {e}")
        return None

def process_data(df):
    """Process the data to extract course code, section, and handle multiple programs"""
    processed_data = []
    
    for _, row in df.iterrows():
        # Extract section (last character of fa_course_section)
        section = str(row['fa_course_section'])[-1] if pd.notna(row['fa_course_section']) else ''
        
        # Extract course code (everything before the first dash)
        course_code = str(row['fa_course_section']).split('-')[0] if pd.notna(row['fa_course_section']) else ''
        
        # Get year from fa_course_year
        year = f"Year {row['fa_course_year']}" if pd.notna(row['fa_course_year']) else 'Year 0'
        
        # Get day and merged timeslot
        day = str(row['fa_day_abbr']) if pd.notna(row['fa_day_abbr']) else ''
        
        # Merge start and end times into timeslot
        start_time = str(row['fa_start_time']) if pd.notna(row['fa_start_time']) else ''
        end_time = str(row['fa_end_time']) if pd.notna(row['fa_end_time']) else ''
        timeslot = f"{start_time} - {end_time}" if start_time and end_time else (start_time or end_time or '')
        
        room = str(row['fa_room_code']) if pd.notna(row['fa_room_code']) else ''
        department = str(row['fa_department']) if pd.notna(row['fa_department']) else ''
        
        # Get program sections and student count
        program_sections = str(row['fa_program_section']) if pd.notna(row['fa_program_section']) else ''
        student_count = int(row['fa_student_count']) if pd.notna(row['fa_student_count']) else 0
        
        # Get the first program name for the Program column (just for display)
        first_program = program_sections.split(',')[0].strip() if program_sections else ''
        program_code = first_program.split('-')[0] if first_program else ''
        program_name = PROGRAM_NAMES.get(program_code, program_code)
        
        processed_data.append({
            'Department': department,
            'Program': program_name,
            'Program Section': program_sections,
            'Student Count': student_count,
            'Year': year,
            'Course Code': course_code,
            'Section': section,
            'Day': day,
            'Timeslot': timeslot,
            'Room': room,
            'fa_program_section': program_sections,
            'fa_student_count': student_count
        })
    
    return pd.DataFrame(processed_data)

def create_excel_file(df):
    """Create Excel workbook with sorted data by Department, Course Code, and Year Level"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    
    # Sort the DataFrame by Department, Course Code, and Year Level
    df_sorted = df.sort_values(by=['Department', 'Course Code', 'Year'])
    
    # Remove duplicate rows (keeping only one entry per course section)
    df_sorted = df_sorted.drop_duplicates(subset=['Department', 'Course Code', 'Section', 'Year', 'Day', 'Timeslot', 'Room'])
    
    current_row = 1
    
    # Get unique departments in sorted order
    departments = sorted(df['Department'].unique())
    
    for department in departments:
        # Write department header
        ws.cell(row=current_row, column=1, value=department)
        current_row += 1
        
        # Get courses for this department in sorted order
        dept_courses = sorted(df_sorted[df_sorted['Department'] == department]['Course Code'].unique())
        
        for course in dept_courses:
            # Write course code header
            ws.cell(row=current_row, column=2, value=course)
            current_row += 1
            
            # Get years for this course in sorted order
            course_years = sorted(df_sorted[(df_sorted['Department'] == department) & 
                                          (df_sorted['Course Code'] == course)]['Year'].unique())
            
            for year in course_years:
                # Write year level header
                ws.cell(row=current_row, column=3, value=year)
                current_row += 1
                
                # Write column headers
                headers = ['', '', 'Section', 'Program Section', 'Student Count', 'Day', 'Timeslot', 'Room']
                for col_num, header in enumerate(headers, start=1):
                    ws.cell(row=current_row, column=col_num, value=header)
                current_row += 1
                
                # Get sections for this department, course and year
                sections = df_sorted[(df_sorted['Department'] == department) & 
                                   (df_sorted['Course Code'] == course) & 
                                   (df_sorted['Year'] == year)]
                
                if sections.empty:
                    # Add empty row if no sections
                    ws.cell(row=current_row, column=3, value='')
                    current_row += 1
                else:
                    # Write section data
                    for _, section in sections.iterrows():
                        ws.cell(row=current_row, column=3, value=section['Section'])
                        ws.cell(row=current_row, column=4, value=section['fa_program_section'])
                        ws.cell(row=current_row, column=5, value=section['fa_student_count'])
                        ws.cell(row=current_row, column=6, value=section['Day'])
                        ws.cell(row=current_row, column=7, value=section['Timeslot'])
                        ws.cell(row=current_row, column=8, value=section['Room'])
                        current_row += 1
                
                # Add empty row after each year
                ws.cell(row=current_row, column=3, value='')
                current_row += 1
    
    # Adjust column widths
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        
        adjusted_width = (max_length + 2) * 1.2
        ws.column_dimensions[column].width = adjusted_width
    
    # Center align all cells
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(horizontal='center')
    
    # Save the Excel file
    excel_file = "Final_Assignments_Sorted.xlsx"
    wb.save(excel_file)
    print(f"Excel file created: {excel_file}")
    return excel_file

def create_word_file(df):
    """Create Word document with sorted data (Times New Roman 12pt, no borders)"""
    doc = Document()
    
    # Set default font to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Sort the DataFrame by Program, Year, and Section (A-Z)
    df_sorted = df.sort_values(by=['Program', 'Year', 'Section'])
    
    # Remove duplicate rows (keeping only one entry per course section)
    df_sorted = df_sorted.drop_duplicates(subset=['Program', 'Course Code', 'Section', 'Year', 'Day', 'Timeslot', 'Room'])
    
    # Get unique programs in sorted order
    programs = sorted(df['Program'].unique())
    
    for program in programs:
        # Add program header (left aligned, bold)
        p = doc.add_paragraph()
        p.add_run(program).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)
        
        # Get years for this program in sorted order
        program_years = sorted(df[df['Program'] == program]['Year'].unique())
        
        for year in program_years:
            # Add year header (left aligned, bold)
            p = doc.add_paragraph()
            p.add_run(year).bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)
            
            # Create table for course data (borderless)
            courses = df_sorted[(df_sorted['Program'] == program) & 
                              (df_sorted['Year'] == year)]
            
            if not courses.empty:
                # Add table (1 row for headers + data rows)
                table = doc.add_table(rows=1, cols=5)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Remove all borders and set font
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                
                # Set header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Course Code'
                hdr_cells[1].text = 'Section'
                hdr_cells[2].text = 'Day'
                hdr_cells[3].text = 'Timeslot'
                hdr_cells[4].text = 'Room'
                
                # Make headers bold and center aligned
                for cell in hdr_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                
                # Add course data
                for _, course in courses.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(course['Course Code']) if pd.notna(course['Course Code']) else ''
                    row_cells[1].text = str(course['Section']) if pd.notna(course['Section']) else ''
                    row_cells[2].text = str(course['Day']) if pd.notna(course['Day']) else ''
                    row_cells[3].text = str(course['Timeslot']) if pd.notna(course['Timeslot']) else ''
                    row_cells[4].text = str(course['Room']) if pd.notna(course['Room']) else ''
                    
                    # Center align all cells
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the Word document
    word_file = "Final_Assignments_Sorted.docx"
    doc.save(word_file)
    print(f"Word document created: {word_file}")
    return word_file

def main():
    # Step 1: Get data from database
    df = get_data_from_db()
    if df is None:
        return
    
    # Step 2: Process the data
    processed_df = process_data(df)
    
    # Step 3: Create Excel file
    excel_file = create_excel_file(processed_df)
    
    # Step 4: Create Word document
    word_file = create_word_file(processed_df)
    
    print(f"\nSuccessfully created both files:")
    print(f"- Excel file: {excel_file}")
    print(f"- Word document: {word_file}")

if __name__ == "__main__":
    main()